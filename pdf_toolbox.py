#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF 工具箱（命令行版）
功能：合并 / 拆分 / 删除页面 / 提取页面 / 调整页序 / 旋转 /
      PDF↔图片 / PDF→Word / Office→PDF / PDF→Excel / PDF→文本
用法：python pdf_toolbox.py <命令> [参数]，详见 python pdf_toolbox.py -h
"""

import argparse
import os
import sys
from pathlib import Path


# ==================== 辅助函数 ====================

def parse_pages(pages_str):
    """把 '1,3,5-7' 解析成 [1, 3, 5, 6, 7]"""
    result = []
    for part in pages_str.split(','):
        part = part.strip()
        if not part:
            continue
        if '-' in part:
            start, end = part.split('-', 1)
            result.extend(range(int(start), int(end) + 1))
        else:
            result.append(int(part))
    return result


def ensure_dir(path):
    """确保输出目录存在"""
    Path(path).mkdir(parents=True, exist_ok=True)


def ensure_extension(path, extension):
    """如果输出路径没有后缀，则自动补充指定后缀。"""
    if not extension.startswith('.'):
        extension = '.' + extension
    output_path = Path(path)
    if output_path.suffix:
        return str(output_path)
    return str(output_path.with_suffix(extension))


def write_pdf(writer, output_path):
    output_path = ensure_extension(output_path, '.pdf')
    writer.write(output_path)
    return output_path


# ==================== 1. 页面组织（pypdf）====================

def cmd_merge(args):
    from pypdf import PdfWriter
    writer = PdfWriter()
    for path in args.inputs:
        print(f"  添加: {path}")
        writer.append(path)
    output = write_pdf(writer, args.output)
    writer.close()
    print(f"✅ 合并完成 -> {output}")


def cmd_split(args):
    from pypdf import PdfReader, PdfWriter
    reader = PdfReader(args.input)
    total = len(reader.pages)
    ensure_dir(args.output_dir)
    stem = Path(args.input).stem
    count = 0
    for start in range(0, total, args.size):
        writer = PdfWriter()
        for i in range(start, min(start + args.size, total)):
            writer.add_page(reader.pages[i])
        count += 1
        out = os.path.join(args.output_dir, f"{stem}_part{count}.pdf")
        writer.write(out)
        print(f"  生成: {out}")
    print(f"✅ 拆分完成，共 {count} 个文件 -> {args.output_dir}")


def cmd_delete(args):
    from pypdf import PdfReader, PdfWriter
    reader = PdfReader(args.input)
    to_delete = set(parse_pages(args.pages))
    writer = PdfWriter()
    for i, page in enumerate(reader.pages, start=1):
        if i not in to_delete:
            writer.add_page(page)
    output = write_pdf(writer, args.output)
    print(f"✅ 已删除页面 {sorted(to_delete)} -> {output}")


def cmd_extract(args):
    from pypdf import PdfReader, PdfWriter
    reader = PdfReader(args.input)
    pages = parse_pages(args.pages)
    writer = PdfWriter()
    for p in pages:
        writer.add_page(reader.pages[p - 1])
    output = write_pdf(writer, args.output)
    print(f"✅ 已提取页面 {pages} -> {output}")


def cmd_reorder(args):
    from pypdf import PdfReader, PdfWriter
    reader = PdfReader(args.input)
    order = parse_pages(args.order)
    writer = PdfWriter()
    for p in order:
        writer.add_page(reader.pages[p - 1])
    output = write_pdf(writer, args.output)
    print(f"✅ 已按新顺序排列 -> {output}")


def cmd_rotate(args):
    from pypdf import PdfReader, PdfWriter
    reader = PdfReader(args.input)
    target = set(parse_pages(args.pages)) if args.pages else None
    writer = PdfWriter()
    for i, page in enumerate(reader.pages, start=1):
        if target is None or i in target:
            page.rotate(args.angle)
        writer.add_page(page)
    output = write_pdf(writer, args.output)
    scope = f"第 {sorted(target)} 页" if target else "全部页面"
    print(f"✅ 已旋转 {scope} {args.angle}° -> {output}")


# ==================== 2. 格式转换 ====================

def cmd_pdf2img(args):
    import fitz  # PyMuPDF
    doc = fitz.open(args.input)
    ensure_dir(args.output_dir)
    stem = Path(args.input).stem
    n = len(doc)
    for i in range(n):
        pix = doc[i].get_pixmap(dpi=args.dpi)
        out = os.path.join(args.output_dir, f"{stem}_p{i + 1}.{args.format}")
        pix.save(out)
        print(f"  生成: {out}")
    doc.close()
    print(f"✅ PDF 转图片完成，共 {n} 页 -> {args.output_dir}")


def cmd_img2pdf(args):
    from PIL import Image
    output = ensure_extension(args.output, '.pdf')
    images = []
    try:
        for p in args.inputs:
            img = Image.open(p)
            images.append(img.convert('RGB') if img.mode != 'RGB' else img.copy())
            img.close()
        images[0].save(output, format='PDF', save_all=True, append_images=images[1:])
        print(f"✅ {len(images)} 张图片合成 PDF -> {output}")
    finally:
        for img in images:
            try:
                img.close()
            except Exception:
                pass


def cmd_pdf2word(args):
    output = ensure_extension(args.output, '.docx')
    if args.mode == 'text':
        pdf2word_text(args.input, output)
        print(f"✅ PDF 转 Word（纯文本优化模式）完成 -> {output}")
        return

    from pdf2docx import Converter
    cv = Converter(args.input)
    try:
        cv.convert(output)
    finally:
        cv.close()
    print(f"✅ PDF 转 Word（版式模式）完成 -> {output}")
    print("ℹ️ 如果版式混乱，可尝试：pdf2word 输入.pdf -o 输出.docx --mode text")


def pdf2word_text(input_pdf, output_docx):
    """把 PDF 文本提取成更稳定、干净的 Word 文档。适合版式转换很乱的 PDF。"""
    import fitz
    from docx import Document
    from docx.shared import Pt

    doc = fitz.open(input_pdf)
    word = Document()
    normal = word.styles['Normal']
    normal.font.name = 'Microsoft YaHei'
    normal.font.size = Pt(10.5)

    for index, page in enumerate(doc, start=1):
        if index > 1:
            word.add_page_break()
        heading = word.add_paragraph()
        heading_run = heading.add_run(f"Page {index}")
        heading_run.bold = True

        text = page.get_text('text').strip()
        if not text:
            word.add_paragraph('[No selectable text found on this page]')
            continue
        for block in text.split('\n'):
            block = block.strip()
            if block:
                word.add_paragraph(block)
    doc.close()
    word.save(output_docx)


def cmd_office2pdf(args):
    import shutil
    import subprocess
    soffice = shutil.which('soffice') or shutil.which('libreoffice')
    if not soffice:
        print("❌ 未找到 LibreOffice，请先安装：https://www.libreoffice.org/")
        return
    ensure_dir(args.output_dir)
    subprocess.run([soffice, '--headless', '--convert-to', 'pdf',
                    '--outdir', args.output_dir, args.input], check=True)
    print(f"✅ Office 转 PDF 完成 -> {args.output_dir}")


def cmd_pdf2excel(args):
    import pdfplumber
    from openpyxl import Workbook
    output = ensure_extension(args.output, '.xlsx')
    wb = Workbook()
    wb.remove(wb.active)
    found = 0
    with pdfplumber.open(args.input) as pdf:
        for pno, page in enumerate(pdf.pages, start=1):
            for tno, table in enumerate(page.extract_tables(), start=1):
                found += 1
                ws = wb.create_sheet(f"P{pno}_T{tno}")
                for row in table:
                    ws.append([("" if c is None else c) for c in row])
    if found == 0:
        print("⚠️ 未在 PDF 中检测到表格")
        return
    wb.save(output)
    print(f"✅ 共提取 {found} 个表格 -> {output}")


def cmd_pdf2text(args):
    import fitz
    output = ensure_extension(args.output, '.txt')
    doc = fitz.open(args.input)
    text = [page.get_text() for page in doc]
    doc.close()
    with open(output, 'w', encoding='utf-8') as f:
        f.write("\n".join(text))
    print(f"✅ 提取文字完成 -> {output}")


# ==================== 命令行定义 ====================

def build_parser():
    parser = argparse.ArgumentParser(description="📄 PDF 工具箱（命令行版）")
    sub = parser.add_subparsers(dest='command', required=True)

    p = sub.add_parser('merge', help='合并多个 PDF')
    p.add_argument('inputs', nargs='+', help='输入 PDF（可多个）')
    p.add_argument('-o', '--output', required=True, help='输出文件，未写 .pdf 会自动补充')
    p.set_defaults(func=cmd_merge)

    p = sub.add_parser('split', help='拆分 PDF')
    p.add_argument('input', help='输入 PDF')
    p.add_argument('-o', '--output-dir', default='output', help='输出目录')
    p.add_argument('-s', '--size', type=int, default=1, help='每个文件页数（默认 1）')
    p.set_defaults(func=cmd_split)

    p = sub.add_parser('delete', help='删除指定页面')
    p.add_argument('input')
    p.add_argument('-p', '--pages', required=True, help='要删除的页码，如 2,4,6-8')
    p.add_argument('-o', '--output', required=True, help='输出文件，未写 .pdf 会自动补充')
    p.set_defaults(func=cmd_delete)

    p = sub.add_parser('extract', help='提取指定页面')
    p.add_argument('input')
    p.add_argument('-p', '--pages', required=True, help='要提取的页码，如 1,3,5-7')
    p.add_argument('-o', '--output', required=True, help='输出文件，未写 .pdf 会自动补充')
    p.set_defaults(func=cmd_extract)

    p = sub.add_parser('reorder', help='调整页面顺序')
    p.add_argument('input')
    p.add_argument('-r', '--order', required=True, help='新顺序，如 3,1,2')
    p.add_argument('-o', '--output', required=True, help='输出文件，未写 .pdf 会自动补充')
    p.set_defaults(func=cmd_reorder)

    p = sub.add_parser('rotate', help='旋转页面')
    p.add_argument('input')
    p.add_argument('-a', '--angle', type=int, choices=[90, 180, 270, -90],
                   required=True, help='旋转角度（顺时针）')
    p.add_argument('-p', '--pages', help='指定页码，不填则全部')
    p.add_argument('-o', '--output', required=True, help='输出文件，未写 .pdf 会自动补充')
    p.set_defaults(func=cmd_rotate)

    p = sub.add_parser('pdf2img', help='PDF 转图片')
    p.add_argument('input')
    p.add_argument('-o', '--output-dir', default='images')
    p.add_argument('--dpi', type=int, default=150, help='清晰度（默认 150）')
    p.add_argument('--format', default='png', choices=['png', 'jpg'])
    p.set_defaults(func=cmd_pdf2img)

    p = sub.add_parser('img2pdf', help='图片转 PDF')
    p.add_argument('inputs', nargs='+', help='输入图片（可多个）')
    p.add_argument('-o', '--output', required=True, help='输出 PDF，未写 .pdf 会自动补充')
    p.set_defaults(func=cmd_img2pdf)

    p = sub.add_parser('pdf2word', help='PDF 转 Word')
    p.add_argument('input')
    p.add_argument('-o', '--output', required=True, help='输出 .docx，未写 .docx 会自动补充')
    p.add_argument('--mode', choices=['layout', 'text'], default='layout',
                   help='layout=尽量保留原版式；text=纯文本优化模式，排版更干净但不保留原版式')
    p.set_defaults(func=cmd_pdf2word)

    p = sub.add_parser('office2pdf', help='Word/PPT 转 PDF（需 LibreOffice）')
    p.add_argument('input', help='输入 .docx/.pptx 等')
    p.add_argument('-o', '--output-dir', default='output')
    p.set_defaults(func=cmd_office2pdf)

    p = sub.add_parser('pdf2excel', help='PDF 表格转 Excel')
    p.add_argument('input')
    p.add_argument('-o', '--output', required=True, help='输出 .xlsx，未写 .xlsx 会自动补充')
    p.set_defaults(func=cmd_pdf2excel)

    p = sub.add_parser('pdf2text', help='PDF 转纯文本')
    p.add_argument('input')
    p.add_argument('-o', '--output', required=True, help='输出 .txt，未写 .txt 会自动补充')
    p.set_defaults(func=cmd_pdf2text)

    return parser


def main():
    args = build_parser().parse_args()
    try:
        args.func(args)
    except Exception as e:
        print(f"❌ 出错了: {e}")
        sys.exit(1)


if __name__ == '__main__':
    main()
